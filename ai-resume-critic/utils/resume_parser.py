"""
Resume Parser - Extracts text from uploaded resume files (PDF, TXT, DOCX).
"""

import io
import re
import streamlit as st


def extract_text_from_upload(uploaded_file) -> str:
    """
    Extract text from an uploaded resume file.
    Supports PDF, TXT, and DOCX formats.
    """
    if uploaded_file is None:
        return ""

    file_type = uploaded_file.type
    file_name = uploaded_file.name.lower()

    try:
        if file_type == "application/pdf" or file_name.endswith(".pdf"):
            return _extract_pdf(uploaded_file)
        elif file_type == "text/plain" or file_name.endswith(".txt"):
            return _extract_txt(uploaded_file)
        elif (
            file_type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            or file_name.endswith(".docx")
        ):
            return _extract_docx(uploaded_file)
        else:
            raise ValueError(
                f"Unsupported file type: {file_type}. "
                "Please upload a PDF, TXT, or DOCX file."
            )
    except Exception as e:
        raise RuntimeError(f"Failed to parse resume: {str(e)}")


def _extract_pdf(uploaded_file) -> str:
    """Extract text from PDF using PyPDF2."""
    try:
        import pypdf
    except ImportError:
        try:
            import PyPDF2 as pypdf
        except ImportError:
            raise ImportError(
                "PDF parsing library not found. "
                "Install with: pip install pypdf"
            )

    try:
        uploaded_file.seek(0)  # reset pointer in case this file was already read on a previous rerun
        pdf_reader = pypdf.PdfReader(io.BytesIO(uploaded_file.read()))
        text_parts = []

        for page_num, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            if page_text and page_text.strip():
                text_parts.append(page_text.strip())

        full_text = "\n".join(text_parts)

        if not full_text.strip():
            raise ValueError(
                "Could not extract text from this PDF. "
                "It might be a scanned document or image-based PDF."
            )

        return _clean_text(full_text)

    except Exception as e:
        if "scanned" in str(e).lower() or "image" in str(e).lower():
            raise ValueError(
                "This PDF appears to be scanned or image-based. "
                "Please upload a text-based PDF or use a TXT/DOCX file."
            )
        raise RuntimeError(f"PDF parsing error: {str(e)}")


def _extract_txt(uploaded_file) -> str:
    """Extract text from plain text file."""
    try:
        uploaded_file.seek(0)  # reset pointer in case this file was already read on a previous rerun
        raw_bytes = uploaded_file.read()
        # Try multiple encodings
        for encoding in ["utf-8", "latin-1", "ascii", "cp1252"]:
            try:
                text = raw_bytes.decode(encoding)
                return _clean_text(text)
            except (UnicodeDecodeError, AttributeError):
                continue
        raise ValueError("Could not decode text file. Please ensure it's a valid text file.")
    except Exception as e:
        raise RuntimeError(f"Text file parsing error: {str(e)}")


def _extract_docx(uploaded_file) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "DOCX parsing library not found. "
            "Install with: pip install python-docx"
        )

    try:
        uploaded_file.seek(0)  # reset pointer in case this file was already read on a previous rerun
        doc = Document(io.BytesIO(uploaded_file.read()))
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text and paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        text_parts.append(cell.text.strip())

        full_text = "\n".join(text_parts)

        if not full_text.strip():
            raise ValueError(
                "Could not extract text from this DOCX file. "
                "The file might be empty or corrupted."
            )

        return _clean_text(full_text)

    except ImportError:
        raise ImportError(
            "DOCX support requires python-docx. "
            "Install with: pip install python-docx"
        )
    except Exception as e:
        raise RuntimeError(f"DOCX parsing error: {str(e)}")


def _clean_text(text: str) -> str:
    """Clean extracted text: remove excessive whitespace, normalize line breaks."""
    if not text:
        return ""

    # Replace multiple spaces with single space
    text = re.sub(r"[^\S\n]+", " ", text)

    # Replace multiple newlines with double newline
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Remove leading/trailing whitespace from each line
    lines = [line.strip() for line in text.split("\n")]
    text = "\n".join(lines)

    return text.strip()


def validate_resume_text(text: str) -> dict:
    """
    Validate extracted resume text and return quality metrics.
    """
    issues = []
    word_count = len(text.split()) if text else 0
    char_count = len(text) if text else 0

    if word_count < 20:
        issues.append("Resume text is extremely short (less than 20 words).")
    elif word_count < 50:
        issues.append("Resume text is very short (less than 50 words).")

    if char_count < 100:
        issues.append("Resume has very little content.")

    # Check for common sections
    common_sections = [
        "experience", "education", "skills", "summary",
        "objective", "projects", "certifications", "achievements",
        "work", "employment", "qualification"
    ]
    text_lower = text.lower()
    found_sections = [s for s in common_sections if s in text_lower]

    if len(found_sections) < 2:
        issues.append(
            f"Only found {len(found_sections)} common resume sections. "
            "A well-structured resume typically has: Experience, Education, Skills, Summary."
        )

    return {
        "word_count": word_count,
        "char_count": char_count,
        "issues": issues,
        "sections_found": found_sections,
        "is_valid": word_count >= 20 and len(found_sections) >= 1,
    }
